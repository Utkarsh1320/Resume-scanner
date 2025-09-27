import io
import os
from docx import Document
import pypdf


async def extract_text_from_file(file_content: bytes, filename: str) -> str:
    file_extension = os.path.splitext(filename)[1].lower()
    print("hello")
    if file_extension == ".pdf":
        return await extract_text_from_pdf(file_content)
    elif file_extension == ".docx":
        return await extract_text_from_docx(file_content)  # You'll need to implement this
    else:
        raise ValueError("Unsupported file type")


async def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            except Exception as e:
                print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                continue
        if not text.strip():
            raise Exception("No text could be extracted from the PDF")
        return text.strip()
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_content)
        document = Document(docx_file)  # read DOCX properly

        text = ""
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"

        if not text.strip():
            raise Exception("No text could be extracted from the DOCX file")

        return text.strip()

    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")